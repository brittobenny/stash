from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

output_path = r"d:\stash\Stash_Test_Cases.docx"
doc = Document()

# Page margins
from docx.shared import Inches
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.underline = (level == 1)
    run.font.size = Pt(13 if level == 1 else 11)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, header_bg="D9D9D9"):
    cols = len(headers)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = str(cell_text)
            for para in row[i].paragraphs:
                para.runs[0].font.size = Pt(9.5) if para.runs else None

    return table


# ─────────────────────────────────────────────────────────────────────
# SECTION 1: TEST CASES
# ─────────────────────────────────────────────────────────────────────
add_heading(doc, "TEST CASES:")

test_case_headers = ["Form / Feature", "Input Data", "Expected Result", "Actual Result", "Remark"]

test_cases = [
    [
        "Register",
        "Valid name, email, mobile number, and password",
        "Create new account and redirect to login page",
        "Account created successfully",
        "Success"
    ],
    [
        "Login (All Users)",
        "Valid email and password",
        "Redirect to corresponding dashboard based on role (Customer / Shop Owner / Admin)",
        "User logged in and redirected to correct dashboard",
        "Success"
    ],
    [
        "Add Pantry Item",
        "Select ingredient from list, enter quantity and optional expiry date",
        "Item added to pantry; quantity updated if already exists",
        "Item saved to pantry successfully",
        "Success"
    ],
    [
        "Update Pantry Quantity",
        "Adjust quantity using +/- controls and save",
        "Pantry item quantity updated in database",
        "Quantity updated successfully",
        "Success"
    ],
    [
        "Expiry Notification",
        "Pantry item with expiry date set within 3 days",
        "System triggers expiry alert notification for the user",
        "Notification displayed correctly",
        "Success"
    ],
    [
        "Low Stock Alert",
        "Pantry item quantity drops below threshold",
        "Low stock warning shown on pantry page; notification created",
        "Alert and notification shown correctly",
        "Success"
    ],
    [
        "Recipe Recommendation",
        "User has pantry items; navigates to Cook with AI",
        "System recommends recipes matching pantry ingredients",
        "Relevant recipes displayed with match percentage",
        "Success"
    ],
    [
        "Add to Cart",
        "Customer selects product and quantity from shop",
        "Product added to cart with correct quantity and price",
        "Cart updated correctly",
        "Success"
    ],
    [
        "Place Order",
        "Customer proceeds to checkout with items in cart",
        "Order created with PLACED status; cart cleared",
        "Order placed and stored in database",
        "Success"
    ],
    [
        "Restock Bill",
        "Customer has low-stock pantry items with matching shop products",
        "Cart auto-filled with matching shop products",
        "Cart populated with low-stock items",
        "Success"
    ],
    [
        "Create Recipe Post",
        "User submits title, ingredients, steps and optional image",
        "Post saved and displayed in social feed",
        "Post created and visible on feed",
        "Success"
    ],
    [
        "Profile Management",
        "User updates name, address, location or profile image",
        "Profile updated across all modules",
        "Changes reflected immediately",
        "Success"
    ],
]

add_table(doc, test_case_headers, test_cases)

# ─────────────────────────────────────────────────────────────────────
# SECTION 2: VALIDATION TESTING
# ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "VALIDATION TESTING")

val_headers = ["Form / Feature", "Input Data", "Expected Result", "Actual Result", "Remark"]

validation_cases = [
    [
        "Register / Login",
        "- Missing email or password\n- Invalid email format\n- Empty fields",
        "Show appropriate validation errors (e.g., 'Email required', 'Invalid format')",
        "Validation messages shown correctly",
        "Success"
    ],
    [
        "Add Pantry Item",
        "- Quantity = 0 or negative\n- No ingredient selected\n- Non-numeric quantity",
        "Error shown: 'quantity must be greater than 0' or 'Please select an ingredient'",
        "Errors triggered as expected",
        "Success"
    ],
    [
        "Update Pantry Quantity",
        "- Quantity set to 0\n- Blank quantity submitted",
        "Error shown: 'Quantity must be greater than 0'",
        "Validation handled properly",
        "Success"
    ],
    [
        "Add Product (Shop Owner)",
        "- Missing product name\n- Price = 0 or negative\n- No category selected",
        "Form prevents submission; error messages shown per field",
        "Errors shown as intended",
        "Success"
    ],
    [
        "Recipe Post",
        "- Empty title\n- No ingredients entered\n- Blank steps",
        "Post prevented; prompt to fill required fields",
        "Validation prevents empty submission",
        "Success"
    ],
    [
        "Recipe Recommendation",
        "- Empty pantry (no items)\n- Pantry with only spices",
        "Show 'No recommendations found' or minimum pantry required message",
        "Handled gracefully with empty state message",
        "Success"
    ],
    [
        "Order / Checkout",
        "- Empty cart\n- Profile incomplete (no address/location)",
        "Error: 'Please complete your profile before checkout'",
        "Validation blocks checkout correctly",
        "Success"
    ],
    [
        "Nutrition Logging (Cook)",
        "- Cooking recipe with 0 quantity\n- Missing serving count",
        "Nutritional values computed as zero or prompt for correction",
        "Graceful handling; no crash",
        "Success"
    ],
]

add_table(doc, val_headers, validation_cases)

# ─────────────────────────────────────────────────────────────────────
# SECTION 3: INTEGRATION TESTING
# ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "INTEGRATION TESTING")

int_headers = ["Form / Feature", "Input Data / Action", "Expected Result", "Actual Result", "Remark"]

integration_cases = [
    [
        "Register",
        "Register a new user with valid details",
        "User should be created in auth_user table and a UserProfile auto-created with role='customer'",
        "Registered successfully; profile created",
        "Success"
    ],
    [
        "Login (All Users)",
        "Enter valid email and password",
        "User validated; token returned; redirected to correct dashboard based on role",
        "Dashboard loaded per role",
        "Success"
    ],
    [
        "Order → Pantry Sync",
        "Customer places order; delivery confirmed by Shop Owner",
        "Order status updated to DELIVERED; pantry quantities incremented automatically",
        "Pantry updated after delivery",
        "Success"
    ],
    [
        "Low Stock → Restock Bill → Cart",
        "Pantry item falls below low stock limit",
        "Low stock alert shown; 'Create Restock Bill' fills cart with matching shop products",
        "Cart auto-filled; checkout flow works",
        "Success"
    ],
    [
        "Recipe Recommendation → Cook Log",
        "User picks recommended recipe and marks it as cooked",
        "CookedRecipeLog created; DailyNutritionScore updated; points awarded",
        "Nutrition and gamification updated correctly",
        "Success"
    ],
    [
        "Shop Product → Pantry Ingredient Link",
        "Shop Owner creates product linked to an Ingredient",
        "On delivery, the product's pack_size is added to the customer's pantry ingredient quantity",
        "Ingredient quantity updated after delivery",
        "Success"
    ],
    [
        "Expiry Detection → Notification",
        "Pantry item's expiry_date is within 3 days",
        "sync_expiry_notifications called on pantry fetch; notification record created for user",
        "Notification visible in notification list",
        "Success"
    ],
    [
        "Recipe Post → Social Feed",
        "User submits a recipe post with title, ingredients and image",
        "Post saved with status='APPROVED'; visible to all users in social feed",
        "Post appears in feed; likes and comments functional",
        "Success"
    ],
    [
        "Admin → User Management",
        "Admin changes a user's role from customer to shopowner",
        "UserProfile.role updated; is_staff flag set on User; user can access shop owner dashboard",
        "Role change reflected across system",
        "Success"
    ],
    [
        "Profile & Access Update",
        "User updates address and location in profile",
        "Changes saved to UserProfile; profile_completed = True; restock bill now accessible",
        "Updated successfully; checkout unblocked",
        "Success"
    ],
]

add_table(doc, int_headers, integration_cases)

doc.save(output_path)
print(f"Test cases document saved to {output_path}")
